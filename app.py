from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Cm
import os
import tkinter as tk
import io # Necesario para guardar la imagen en memoria


class DocCreatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Crear documento de Word con imágenes")

        self.uploaded_images_data = []
        self.image_display_frames = []
        self.dimensions_entries = [] # Stores {"width": Entry_Widget, "height": Entry_Widget, "original_path": str}

        # --- Nueva sección de control global de dimensiones ---
        self.global_dimensions_frame = tk.LabelFrame(self.root, text="Dimensiones globales para todas las imágenes")
        self.global_dimensions_frame.pack(pady=10, padx=10, fill=tk.X) # Pack this frame at the top

        tk.Label(self.global_dimensions_frame, text="Ancho (cm):").grid(row=0, column=0, padx=5, pady=5)
        self.global_width_entry = tk.Entry(self.global_dimensions_frame, width=10)
        self.global_width_entry.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(self.global_dimensions_frame, text="Altura (cm):").grid(row=0, column=2, padx=5, pady=5)
        self.global_height_entry = tk.Entry(self.global_dimensions_frame, width=10)
        self.global_height_entry.grid(row=0, column=3, padx=5, pady=5)

        self.apply_all_button = tk.Button(self.global_dimensions_frame, text="Aplicar a todas", command=self.apply_global_dimensions)
        self.apply_all_button.grid(row=0, column=4, padx=10, pady=5)
        self.apply_all_button.config(state=tk.DISABLED) # Deshabilitar inicialmente

        # --- Fin de la nueva sección de control global ---

        self.upload_button = tk.Button(
            self.root, text="Subir varias imágenes", command=self.upload_images
        )
        self.upload_button.pack(pady=10)

        self.create_doc_button = tk.Button(
            self.root,
            text="Crear documento de Word",
            command=self.create_document_with_images,
        )
        self.create_doc_button.pack(pady=5)
        self.create_doc_button.config(state=tk.DISABLED)

        # Usar un Frame para contener el Canvas y Scrollbar para mejor organización
        canvas_frame = tk.Frame(self.root)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.canvas = tk.Canvas(canvas_frame)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.scrollbar = tk.Scrollbar(
            canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview
        )
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")),
        )

        self.images_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.images_frame, anchor="nw")

        self.images_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion = self.canvas.bbox("all")))
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_mousewheel)
        self.canvas.bind_all("<Button-5>", self._on_mousewheel)

        self.columns_per_row = 5 # Puedes ajustar este número

    def _on_mousewheel(self, event):
        if self.canvas.winfo_exists():
            if event.num == 4 or event.delta > 0:
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5 or event.delta < 0:
                self.canvas.yview_scroll(1, "units")

    def upload_images(self):
        file_paths = filedialog.askopenfilenames(
            title="Escoger imágenes", filetypes=[("Image files", "*.png *.jpg *.jpeg")]
        )

        if not file_paths:
            return

        for frame in self.image_display_frames:
            frame.destroy()
        self.image_display_frames.clear()
        self.uploaded_images_data.clear()
        self.dimensions_entries.clear()

        for i, file_path in enumerate(file_paths):
            try:
                image = Image.open(file_path)
                self.uploaded_images_data.append({"path": file_path, "image": image})
                self.display_image(
                    image, i, file_path
                )
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Error procesando imagen {os.path.basename(file_path)}: {e}",
                )

        self.images_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

        if self.uploaded_images_data:
            self.create_doc_button.config(state=tk.NORMAL)
            self.apply_all_button.config(state=tk.NORMAL) # Habilitar el botón "Aplicar a todas"
        else:
            self.create_doc_button.config(state=tk.DISABLED)
            self.apply_all_button.config(state=tk.DISABLED) # Deshabilitar si no hay imágenes

    def display_image(self, image, index, original_path):
        row = index // self.columns_per_row
        col = index % self.columns_per_row

        image_frame = tk.LabelFrame(self.images_frame, text=f"Imagen {index+1}")
        image_frame.grid(row=row, column=col, padx=10, pady=10, sticky="n")
        self.image_display_frames.append(image_frame)

        img_width, img_height = image.size
        max_size = 200
        if img_width > max_size or img_height > max_size:
            ratio = min(max_size / img_width, max_size / img_height)
            image = image.resize(
                (int(img_width * ratio), int(img_height * ratio)), Image.LANCZOS
            )

        photo = ImageTk.PhotoImage(image)
        image_label = tk.Label(image_frame, image=photo)
        image_label.image = photo
        image_label.pack(pady=5)

        dimensions_frame = tk.Frame(image_frame)
        dimensions_frame.pack(pady=5)

        tk.Label(dimensions_frame, text="Ancho (cm):").grid(row=0, column=0)
        width_entry = tk.Entry(dimensions_frame, width=10)
        width_entry.grid(row=0, column=1)
        width_cm = round(img_width * (2.54 / 96), 2)
        width_entry.insert(0, str(width_cm))

        tk.Label(dimensions_frame, text="Altura (cm):").grid(row=1, column=0)
        height_entry = tk.Entry(dimensions_frame, width=10)
        height_entry.grid(row=1, column=1)
        height_cm = round(img_height * (2.54 / 96), 2)
        height_entry.insert(0, str(height_cm))

        self.dimensions_entries.append(
            {
                "width": width_entry,
                "height": height_entry,
                "original_path": original_path,
            }
        )

    def apply_global_dimensions(self):
        """Aplica las dimensiones globales a todos los campos de entrada de las imágenes."""
        try:
            global_width = float(self.global_width_entry.get())
            global_height = float(self.global_height_entry.get())

            if global_width <= 0 or global_height <= 0:
                messagebox.showwarning("Entrada inválida", "El ancho y la altura deben ser valores positivos.")
                return

            for entry_set in self.dimensions_entries:
                entry_set["width"].delete(0, tk.END)
                entry_set["width"].insert(0, str(global_width))
                entry_set["height"].delete(0, tk.END)
                entry_set["height"].insert(0, str(global_height))
            
            messagebox.showinfo("Aplicado", "Dimensiones aplicadas a todas las imágenes.")

        except ValueError:
            messagebox.showerror("Error de entrada", "Por favor, ingrese valores numéricos válidos para el ancho y la altura globales.")
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error al aplicar las dimensiones: {e}")


    def create_document_with_images(self):
        images_to_insert = []
        for entry_set in self.dimensions_entries:
            try:
                width_cm = float(entry_set["width"].get())
                height_cm = float(entry_set["height"].get())
                original_path = entry_set["original_path"]
                images_to_insert.append((original_path, width_cm, height_cm))
            except ValueError:
                messagebox.showerror(
                    "Error de entrada",
                    f"Dimensiones inválidas para imagen {os.path.basename(entry_set['original_path'])}. Ingrese solo números.",
                )
                return

        if not images_to_insert:
            messagebox.showinfo(
                "Sin imágenes", "No hay imágenes con dimensiones válidas para crear el documento."
            )
            return

        doc_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos de Word", "*.docx")],
            title="Guardar documento de Word",
        )

        if not doc_path:
            return

        try:
            inserted_count, error_count = self.insert_image_with_size(doc_path, images_to_insert)
            messagebox.showinfo("Éxito", f"{inserted_count} imágenes insertadas en el documento guardado en {doc_path}. {error_count} imágenes no se pudieron insertar.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo crear el documento: {e}")

    def insert_image_with_size(self, doc_path, images_data):
        document = Document()
        inserted_count = 0
        error_count = 0

        for image_path, width_cm, height_cm in images_data:
            try:
                # Abrir la imagen con Pillow
                img_pil = Image.open(image_path)

                # Convertir a RGB si es necesario (para evitar problemas con modos como RGBA o P)
                if img_pil.mode in ("RGBA", "P"):
                    img_pil = img_pil.convert("RGB")

                # Guardar la imagen en un buffer de memoria como PNG para docx
                img_byte_arr = io.BytesIO()
                img_pil.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0) # Regresar al inicio del buffer

                document.add_paragraph() # Esto crea un nuevo párrafo para cada imagen
                paragraph = document.paragraphs[0]
                paragraph.add_run().add_picture(
                        img_byte_arr, width=Cm(width_cm), height=Cm(height_cm)
                    )
                    # Add image name and dimensions below the image
                paragraph.add_run("  ")  # Add an empty paragraph for spacing between images
                inserted_count += 1

            except Exception as e:
                error_count += 1
                continue

        document.save(doc_path)
        return inserted_count, error_count


if __name__ == "__main__":
    root = tk.Tk()
    app = DocCreatorApp(root)
    root.geometry("800x800") # Aumentar un poco el tamaño inicial de la ventana
    root.mainloop()