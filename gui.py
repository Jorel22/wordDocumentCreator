from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Cm
import os
import tkinter as tk


class DocCreatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Crear documento de Word con imágenes")

        self.uploaded_images_data = []
        self.image_display_frames = []
        self.dimensions_entries = []
        self.upload_button = tk.Button(
            self.root, text="Subir varias imágenes", command=self.upload_images
        )
        self.upload_button.pack(pady=10)

        # Button to create the Word document
        self.create_doc_button = tk.Button(
            self.root,
            text="Crear documento de Word",
            command=self.create_document_with_images,
        )
        self.create_doc_button.pack(pady=5)
        self.create_doc_button.config(state=tk.DISABLED)  # Disable initially

        self.canvas = tk.Canvas(root)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.scrollbar = tk.Scrollbar(
            root, orient=tk.VERTICAL, command=self.canvas.yview
        )
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")),
        )

        self.images_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.images_frame, anchor="nw")

        # Bind scroll wheel to the canvas (for better UX)
        self.images_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion = self.canvas.bbox("all")))
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel) # Windows/macOS
        self.canvas.bind_all("<Button-4>", self._on_mousewheel) # Linux scroll up
        self.canvas.bind_all("<Button-5>", self._on_mousewheel) # Linux scroll down

        # Define how many columns you want per row
        self.columns_per_row = 5 # You can adjust this number
    
    def _on_mousewheel(self, event):
        if self.canvas.winfo_exists(): # Check if canvas still exists
            if event.num == 4 or event.delta > 0: # Scroll up
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5 or event.delta < 0: # Scroll down
                self.canvas.yview_scroll(1, "units")

    def upload_images(self):
        file_paths = filedialog.askopenfilenames(
            title="Escoger imágenes", filetypes=[("Image files", "*.png *.jpg *.jpeg")]
        )

        if not file_paths:
            return

        # Clear previous images and data
        for frame in self.image_display_frames:
            frame.destroy()
        self.image_display_frames.clear()
        self.uploaded_images_data.clear()
        self.dimensions_entries.clear()

        for i, file_path in enumerate(file_paths):
            try:
                image = Image.open(file_path)
                self.uploaded_images_data.append({"path": file_path, "image": image})
                self.display_image(
                    image, i, file_path
                )  # Pass file_path to display_image
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Error procesando imagen {os.path.basename(file_path)}: {e}",
                )
                
        self.images_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

        # Enable the "Create Doc" button if images are uploaded
        if self.uploaded_images_data:
            self.create_doc_button.config(state=tk.NORMAL)
        else:
            self.create_doc_button.config(state=tk.DISABLED)

    def display_image(self, image, index, original_path):
        row = index // self.columns_per_row
        col = index % self.columns_per_row

        image_frame = tk.LabelFrame(self.images_frame, text=f"Imagen {index+1}")
        image_frame.grid(row=row, column=col, padx=10, pady=10, sticky="n")
        self.image_display_frames.append(image_frame)

        # Resize image for display to fit within a reasonable size
        img_width, img_height = image.size
        max_size = 200
        if img_width > max_size or img_height > max_size:
            ratio = min(max_size / img_width, max_size / img_height)
            image = image.resize(
                (int(img_width * ratio), int(img_height * ratio)), Image.LANCZOS
            )

        photo = ImageTk.PhotoImage(image)
        image_label = tk.Label(image_frame, image=photo)
        image_label.image = photo  # Keep a reference to prevent garbage collection
        image_label.pack(pady=5)

        # Dimensions input
        dimensions_frame = tk.Frame(image_frame)
        dimensions_frame.pack(pady=5)

        tk.Label(dimensions_frame, text="Ancho (cm):").grid(row=0, column=0)
        width_entry = tk.Entry(dimensions_frame, width=10)
        width_entry.grid(row=0, column=1)
        # Convert pixels to cm for initial display (approx. 96 dpi)
        # 1 inch = 2.54 cm, 1 inch = 96 pixels (common assumption for screen)
        # So, 1 pixel = 2.54 / 96 cm
        width_cm = round(img_width * (2.54 / 96), 2)
        width_entry.insert(0, str(width_cm))

        tk.Label(dimensions_frame, text="Altura (cm):").grid(row=1, column=0)
        height_entry = tk.Entry(dimensions_frame, width=10)
        height_entry.grid(row=1, column=1)
        height_cm = round(img_height * (2.54 / 96), 2)
        height_entry.insert(0, str(height_cm))

        self.dimensions_entries.append(
            {
                "width": width_entry,
                "height": height_entry,
                "original_path": original_path,
            }
        )

    def create_document_with_images(self):
        # Gather dimensions and paths from the input fields
        images_to_insert = []  # List of tuples: (image_path, width_cm, height_cm)
        for entry_set in self.dimensions_entries:
            try:
                width_cm = float(entry_set["width"].get())
                height_cm = float(entry_set["height"].get())
                original_path = entry_set["original_path"]
                images_to_insert.append((original_path, width_cm, height_cm))
            except ValueError:
                messagebox.showerror(
                    "Input Error",
                    f"Dimensiones inválidas para imagen {os.path.basename(entry_set['original_path'])}. Ingrese solo números.",
                )
                return

        if not images_to_insert:
            messagebox.showinfo(
                "No Images", "No images with valid dimensions to create the document."
            )
            return

        # Ask the user for the save path of the .docx file
        doc_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Guardar documento de Word",
        )

        if not doc_path:
            return  # User cancelled

        try:
            inserted_images, images_with_error = self.insert_image_with_size(doc_path, images_to_insert)
            messagebox.showinfo("Éxito", f" {inserted_images} imágenes insertadas en documento guardado en {doc_path}, {images_with_error} imágenes no insertadas")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo crear el document: {e}")

    def insert_image_with_size(self, doc_path, images_data):
            document = Document()
            correct_images = 0
            for image_path, width_cm, height_cm in images_data:
                try:
                    document.add_paragraph()  # Add a paragraph to insert the image into
                    paragraph = document.paragraphs[0]
                    paragraph.add_run().add_picture(
                        image_path, width=Cm(width_cm), height=Cm(height_cm)
                    )
                    # Add image name and dimensions below the image
                    paragraph.add_run("  ")  # Add an empty paragraph for spacing between images
                    correct_images += 1
                except Exception as e:
                    continue  # Skip this image and continue with the next one
            document.save(doc_path)
            incorrect_images = len(images_data) - correct_images
            return correct_images, incorrect_images
            # raise e


if __name__ == "__main__":
    root = tk.Tk()
    app = DocCreatorApp(root)
    root.geometry("800x600")  # Set initial window size
    root.mainloop()
